from pathlib import Path
from typing import List

from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document


class CustomDocxLoader(BaseLoader):
    """
    LangChain-compatible loader for .docx files.

    langchain_community ships Docx2txtLoader / UnstructuredWordDocumentLoader,
    but neither pulls table content out cleanly, and evidence documents
    (charge-sheets, structured witness-statement forms) frequently carry
    key facts in tables. This custom loader keeps that table-aware
    extraction while conforming to the standard BaseLoader interface, so
    it drops into DocumentLoaderFactory exactly like a community loader
    would.
    """

    def __init__(self, file_path: str):
        self._file_path = file_path

    def load(self) -> List[Document]:
        import docx  # python-docx

        document = docx.Document(self._file_path)

        paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]

        table_lines = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))

        text = "\n".join(paragraphs + table_lines)
        filename = Path(self._file_path).name

        return [
            Document(
                page_content=text,
                metadata={"source": self._file_path, "filename": filename, "page": 0},
            )
        ]
